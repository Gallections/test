from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn
import pdfplumber
# import pytesseract
import re
from PIL import Image

def extract_text_within_bbox(page, bbox):
    cropped_page = page.within_bbox(bbox)
    text = cropped_page.extract_text()
    
    # If no text is found, use pytesseract to extract it
    if not text:
        image = cropped_page.to_image(resolution=300)
        text = pytesseract.image_to_string(image.image)

    return text

# Function to join lines based on specific conditions
def join_lines(text):
    lines = text.split("\n")
    new_lines = []
    for i in range(len(lines)):
        if i > 0 and (lines[i].startswith('(') or lines[i][0].islower()):
            new_lines[-1] = new_lines[-1].strip() + ' ' + lines[i]
        elif i < len(lines) - 1 and (lines[i].endswith(',') or lines[i][-1].isdigit()):
            lines[i+1] = lines[i] + ' ' + lines[i+1]
        else:
            new_lines.append(lines[i])
    return "\n".join(new_lines)

# adds dashes before an element in list gets indented
def add_dashes(sample_list):
    new_list = []
    for element in sample_list:
        new_list.append( element.replace("\n", "\n- "))
    return new_list 
    

# Function to split text into subgroups based on a regex pattern
def split_into_subgroups(text):
    text = join_lines(text)
    new_text = text.split("\n")
    subgroups = []
    counter = -1
    pattern = r'^(ECP|Avantaguard|Appearance|Protection)'

    for line in new_text:
        if re.match(pattern, line):
            subgroups.append("")
            counter += 1
        if counter >= 0:  # Ensure there's at least one subgroup
            subgroups[counter] += line + "\n"

    return [group.strip() for group in subgroups if group.strip()]

# Removes empty strings from the text
def remove_empty_strings(string_list):
    return list(filter(None, string_list))

# Converts the text(s) to a list
def text_to_list(*args):
    res = []
    for arg in args:
        res.append(arg)
    return res

# doc.add_paragraph("Hello, World!")
doc = Document()

# Replace every key terms with list of data.
def fill_info(template, output, placeholder):
    doc = Document(template)
    data_counter = 0
    for paragraph in doc.paragraphs:
        while paragraph.text.find(placeholder) != -1:
            paragraph.text = paragraph.text.replace(placeholder, str(data[data_counter]), 1)
            data_counter += 1
            
    doc.save(output)

# testing data dictionaries
info_dict = {"Client": ["Name"], 
             "Vehicle": ["2024 Toyota bX4X... "], 
             "Financial information": ["Random studff , 3.99% ..... random information again"]}

other_dict = {"Premium": ["one", "two", "three", "five"], 
              "Preferred": ["one", "TWO", "THREE", "four", "five"],
              "Essential": ["ONE", "TSO"]}

payment_dict = {"Approved Biweekly Payment": ["193.02", "188.89", "182.78"],
                "Client wishes to take": []}

# finds the max size of dictionary value
def longest_values_size(sample):
    max_size = 0
    for value in sample.values():
        max_size = max(max_size, len(value))
    return max_size + 1

# adds the main Toyota logo
def add_main_logo(logo, output):
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run()
    inline_shape = run.add_picture(logo) 
    inline_shape.width = Inches(1.9)
    inline_shape.height = Inches(1.7)

    doc.save(output)

# adds the bottom sponsor logos
def add_sponsors(output, *logos):
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    for logo in logos:
        run = paragraph.add_run()
        inline_shape = run.add_picture(logo) 
        # inline_shape.width = Inches(1)
        # inline_shape.height = Inches(0.8)
    doc.save(output)

def add_line_breaks(output):
    paragraph = doc.add_paragraph("       ")
    run = paragraph.runs[0]

    # Set the font size
    font_size = Pt(3)  # Specify the font size in points
    run.font.size = font_size

    doc.save(output)


# Filling in the Client section of the table
def fill_client_dict(output, data):
    row_size = len(data)
    col_size = longest_values_size(data)
    doc.add_table(rows= row_size, cols = col_size)

    table = doc.tables[-1]

    keys = list(data.keys())

    for i in range(row_size):
        values = list(data.values())[i]
        for j in range(col_size):
            if j == 0:
                cell = table.rows[i].cells[j]
                cell.text = keys[i]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                cell_xml_element = cell._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn("w:fill"), "#9fc0f5")
                table_cell_properties.append(shade_obj)
            else:
                table.rows[i].cells[j].text = (values[j - 1])
                cell_xml_element = table.rows[i].cells[j]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn("w:fill"), "#9fc0f5")
                table_cell_properties.append(shade_obj)
    table.style = 'Table Grid'
    doc.save(output)

# filling in the middle section of the table
def fill_preferrence_dict(output, data):
    col_size = len(data) + 1
    row_size = longest_values_size(data)
    doc.add_table(rows= row_size, cols = col_size)

    table = doc.tables[-1]
    keys = list(data.keys())

    for i in range(col_size):
        for j in range(row_size):
            if i == 0:
                continue
            elif j == 0:
                cell = table.columns[i].cells[j]
                cell.text = keys[i -1]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            # elif i == 0 and j == (row_size - 1):
            #     table.columns[i].cells[j].text = "Approved Biweekly Payment"
            # elif i == 0 and j == (row_size - 2):
            #     table.columns[i].cells[j].text = "Client wishes to take"
            else:
                values = list(data.values())[i - 1]
                if j <= len(values):
                    cell = table.columns[i].cells[j]
                    # cell.text = values[j - 1]

                    cell_text = values[j - 1]
                    indent = cell_text.find("\n")
                    if indent == -1:
                        indent = len(cell_text)
                    cell_paragraph = cell.paragraphs[0]
                    cell_paragraph.add_run(cell_text[:indent]).bold = True  # Add the non-bold part
                    cell_paragraph.add_run(cell_text[indent:]).font.size = Pt(10) # Add the bold part


    table.style = 'Table Grid'
    doc.save(output)

# Filling in the payment section of the table
def fill_payment_dict (output, data):
    row_size = len(data)
    col_size = longest_values_size(data)
    doc.add_table(rows= row_size, cols = col_size)

    table = doc.tables[-1]

    keys = list(data.keys())

    for i in range(row_size):
        values = list(data.values())[i]
        for j in range(col_size):
            if j == 0:
                cell = table.rows[i].cells[j]
                cell.text = keys[i]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                cell_xml_element = cell._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn("w:fill"), "#9fc0f5")
                table_cell_properties.append(shade_obj)
                
            else:
                if j <= len(values):
                    table.rows[i].cells[j].text = str(values[j - 1])
                cell_xml_element = table.rows[i].cells[j]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn("w:fill"), "#9fc0f5")
                table_cell_properties.append(shade_obj)
    table.style = 'Table Grid'
    doc.save(output)

def do_everything(file) :
    # Open the PDF and read the first page
    with pdfplumber.open(file) as pdf:
        first_page = pdf.pages[0]
        
        # Define bounding boxes for each section
        name_bbox = (148, 23, 333, 38)
        model_bbox = (148, 37, 332, 59)
        premium_bbox = (28, 75, 268, 513)
        preferred_bbox = (276, 75, 516, 513)
        essential_bbox = (523, 75, 763, 513)
        premium_price_bbox = (28, 513, 270, 528)
        preferred_price_bbox = (276, 513, 516, 528)
        essential_price_bbox = (523, 513, 763, 528)
        
        # Extract text under each section using the defined bounding boxes also name and model
        name_text = extract_text_within_bbox(first_page, name_bbox)
        model_text = extract_text_within_bbox(first_page, model_bbox)
        premium_text = extract_text_within_bbox(first_page, premium_bbox)
        preferred_text = extract_text_within_bbox(first_page, preferred_bbox)
        essential_text = extract_text_within_bbox(first_page, essential_bbox)
        premium_price_text = extract_text_within_bbox(first_page, premium_price_bbox)
        preferred_price_text = extract_text_within_bbox(first_page, preferred_price_bbox)
        essential_price_text = extract_text_within_bbox(first_page, essential_price_bbox)

    # Split extracted texts into subgroups using key terms
    premium_list = add_dashes(remove_empty_strings(split_into_subgroups(premium_text)))
    preferred_list = add_dashes(remove_empty_strings(split_into_subgroups(preferred_text)))
    essential_list = add_dashes(remove_empty_strings(split_into_subgroups(essential_text)))

    client_list = text_to_list(name_text)
    print(client_list)
    print()

    vehicle_list = text_to_list(model_text)
    print(vehicle_list)
    print()

    finance_list = text_to_list("")

    biweekly_list = text_to_list(premium_price_text, preferred_price_text, essential_price_text)

    client_wishes_take_list = []

    client_info_dict = {"Client": client_list, 
                        "Vehicle": vehicle_list, 
                        "Financial Information": finance_list}
    preferrence_info_dict = {"Premium": premium_list,
                            "Preferred": preferred_list,
                            "Essential": essential_list}

    payment_info_dict = {"Biweekly Payment": biweekly_list, 
                        "Client wishes to take": client_wishes_take_list}

    add_main_logo("logos/Toyota.png", "table_template.docx")
    add_line_breaks("table_template.docx")
    fill_client_dict("table_template.docx", client_info_dict)
    fill_preferrence_dict("table_template.docx", preferrence_info_dict)
    fill_payment_dict("table_template.docx", payment_info_dict)
    add_sponsors("table_template.docx", "logos/RBC.png", "logos/CIBC.png", "logos/GBC.png", "logos/ScotiaBank.png", "logos/National_Bank.png", "logos/TD.png")


# fill_info("template.docx", "template2.docx", "[]")